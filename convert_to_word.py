#!/usr/bin/env python3
"""
Script to convert UTB HR Central User Manual from Markdown to Word Document
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

def markdown_to_word(md_file, docx_file):
    """Convert Markdown file to Word document"""
    
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    lines = content.split('\n')
    i = 0
    in_code_block = False
    code_block_content = []
    
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        
        # Handle code blocks
        if stripped.startswith('```'):
            if in_code_block:
                # End code block
                if code_block_content:
                    p = doc.add_paragraph('\n'.join(code_block_content))
                    p.style = 'No Spacing'
                    for run in p.runs:
                        run.font.name = 'Courier New'
                        run.font.size = Pt(9)
                code_block_content = []
                in_code_block = False
            else:
                in_code_block = True
            i += 1
            continue
        
        if in_code_block:
            code_block_content.append(line)
            i += 1
            continue
        
        # Skip empty lines
        if not stripped:
            i += 1
            continue
        
        # Title (single #)
        if stripped.startswith('# '):
            title = stripped[2:].strip()
            p = doc.add_heading(title, level=1)
            i += 1
        
        # Heading level 2
        elif stripped.startswith('## '):
            heading = stripped[3:].strip()
            p = doc.add_heading(heading, level=2)
            i += 1
        
        # Heading level 3
        elif stripped.startswith('### '):
            heading = stripped[4:].strip()
            p = doc.add_heading(heading, level=3)
            i += 1
        
        # Heading level 4
        elif stripped.startswith('#### '):
            heading = stripped[5:].strip()
            p = doc.add_heading(heading, level=4)
            i += 1
        
        # Horizontal rule
        elif stripped.startswith('---'):
            i += 1
            continue
        
        # Table detection
        elif '|' in stripped and i + 1 < len(lines):
            # Check if next line is separator
            next_line = lines[i + 1].strip() if i + 1 < len(lines) else ''
            if '|' in next_line and '---' in next_line:
                # Parse table
                table_lines = []
                j = i
                while j < len(lines) and '|' in lines[j].strip():
                    if '---' not in lines[j]:
                        table_lines.append(lines[j])
                    j += 1
                
                if len(table_lines) >= 1:
                    # Create table
                    headers = [cell.strip() for cell in table_lines[0].split('|')[1:-1]]
                    table = doc.add_table(rows=1, cols=len(headers))
                    table.style = 'Light Grid Accent 1'
                    
                    # Add headers
                    header_cells = table.rows[0].cells
                    for k, header in enumerate(headers):
                        header_cells[k].text = header
                        header_cells[k].paragraphs[0].runs[0].bold = True
                    
                    # Add data rows
                    for row_line in table_lines[1:]:
                        row_data = [cell.strip() for cell in row_line.split('|')[1:-1]]
                        if row_data:
                            row = table.add_row()
                            for k, cell_data in enumerate(row_data):
                                if k < len(row.cells):
                                    row.cells[k].text = cell_data
                
                i = j
                continue
        
        # Bullet list
        elif stripped.startswith('- ') or stripped.startswith('* '):
            items = []
            j = i
            while j < len(lines) and (lines[j].strip().startswith('- ') or lines[j].strip().startswith('* ')):
                item_text = lines[j].strip()[2:].strip()
                # Process inline formatting
                item_text = process_inline_formatting(item_text)
                items.append(item_text)
                j += 1
            
            for item in items:
                p = doc.add_paragraph(item, style='List Bullet')
            
            i = j
            continue
        
        # Numbered list
        elif re.match(r'^\d+\.\s', stripped):
            items = []
            j = i
            while j < len(lines) and re.match(r'^\d+\.\s', lines[j].strip()):
                item_text = re.sub(r'^\d+\.\s', '', lines[j].strip())
                item_text = process_inline_formatting(item_text)
                items.append(item_text)
                j += 1
            
            for item in items:
                p = doc.add_paragraph(item, style='List Number')
            
            i = j
            continue
        
        # Regular paragraph
        else:
            text = process_inline_formatting(stripped)
            p = doc.add_paragraph(text)
            i += 1
    
    # Save document
    doc.save(docx_file)
    print(f"Successfully converted {md_file} to {docx_file}")

def process_inline_formatting(text):
    """Process inline markdown formatting"""
    # Remove markdown formatting but keep text
    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)  # Bold
    text = re.sub(r'\*(.*?)\*', r'\1', text)  # Italic
    text = re.sub(r'`(.*?)`', r'\1', text)  # Code
    text = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', text)  # Links
    return text

if __name__ == '__main__':
    import sys
    
    md_file = 'UTB_HR_Central_User_Manual.md'
    docx_file = 'UTB_HR_Central_User_Manual.docx'
    
    if len(sys.argv) > 1:
        md_file = sys.argv[1]
    if len(sys.argv) > 2:
        docx_file = sys.argv[2]
    
    try:
        markdown_to_word(md_file, docx_file)
        print(f"\nWord document created: {docx_file}")
    except Exception as e:
        print(f"Error: {e}")
        import traceback
        traceback.print_exc()
